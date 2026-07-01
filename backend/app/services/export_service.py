import io
import logging
import re

logger = logging.getLogger(__name__)

# 회의록에 저장되는 섹션 색 마커: <!-- section-color:#00754a -->
_COLOR_LINE = re.compile(r"^<!--\s*section-color:\s*(#[0-9a-fA-F]{3,8})\s*-->\s*$")
_COLOR_ANY = re.compile(r"<!--\s*section-color:[^>]*-->\s*\n?")


def _hex_to_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


class ExportService:
    """Exports meeting notes to MD, PDF, or DOCX (섹션 색 반영)."""

    async def export(self, notes_md: str, title: str, fmt: str) -> tuple[bytes, str]:
        if fmt == "md":
            # 색 마커는 UI 전용 → 다운로드 .md 에선 제거해 깔끔하게
            clean = _COLOR_ANY.sub("", notes_md)
            return clean.encode("utf-8"), "text/markdown"
        elif fmt == "pdf":
            return self._to_pdf(notes_md, title), "application/pdf"
        elif fmt == "docx":
            return self._to_docx(notes_md, title), (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            raise ValueError(f"Unsupported format: {fmt}")

    def _to_pdf(self, md_text: str, title: str) -> bytes:
        import markdown
        from weasyprint import HTML

        # 색 마커 + 다음 헤딩 → 색상 스타일 헤딩(HTML)으로 치환
        def repl(m: re.Match) -> str:
            color, hashes, text = m.group(1), m.group(2), m.group(3)
            lvl = len(hashes)
            return (
                f'<h{lvl} class="sec" style="color:{color};'
                f'border-left:5px solid {color};background:{color}14">{text}</h{lvl}>'
            )

        pre = re.sub(
            r"<!--\s*section-color:\s*(#[0-9a-fA-F]{3,8})\s*-->\s*\n\s*(#{1,4})\s+(.+)",
            repl,
            md_text,
        )
        pre = _COLOR_ANY.sub("", pre)  # 남은 색 주석 제거
        html_body = markdown.markdown(pre, extensions=["tables"])

        html_full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
@page {{ margin: 1.8cm; }}
body {{ font-family: "Noto Sans CJK KR", "Noto Sans KR", sans-serif;
        color: #1c1b18; line-height: 1.65; font-size: 10.5pt; }}
h1 {{ font-size: 1.75em; margin: 0 0 .25em; padding-bottom: .3em;
      border-bottom: 2px solid #00754a; }}
h2 {{ font-size: 1.2em; margin: 1.3em 0 .5em; color: #00754a; }}
h2.sec, h3.sec {{ padding: .32em .7em; border-radius: 6px; }}
h3 {{ font-size: 1.05em; margin: 1em 0 .3em; color: #33433d; }}
p {{ margin: .4em 0; }}
strong {{ color: #111; }}
ul, ol {{ padding-left: 1.3em; margin: .4em 0; }}
li {{ margin: .25em 0; }}
table {{ border-collapse: collapse; width: 100%; margin: .7em 0; font-size: .95em; }}
th, td {{ border: 1px solid #e2ded4; padding: 7px 10px; text-align: left; }}
th {{ background: #f2f0eb; color: #00754a; }}
</style></head>
<body>{html_body}</body></html>"""
        return HTML(string=html_full).write_pdf()

    def _to_docx(self, md_text: str, title: str) -> bytes:
        from docx import Document
        from docx.shared import RGBColor

        doc = Document()
        doc.add_heading(title, 0)

        pending_color: str | None = None
        for raw in md_text.split("\n"):
            s = raw.strip()
            if not s:
                continue

            cm = _COLOR_LINE.match(s)
            if cm:
                pending_color = cm.group(1)  # 다음 헤딩에 적용
                continue

            if s.startswith(("# ", "## ", "### ")):
                level = 1 if s.startswith("# ") else 2 if s.startswith("## ") else 3
                h = doc.add_heading(s[level + 1:], level=level)
                color = pending_color or "#00754a"  # 색 없으면 브랜드 그린
                rgb = RGBColor(*_hex_to_rgb(color))
                for run in h.runs:
                    run.font.color.rgb = rgb
                pending_color = None
            elif s.startswith(("- [ ] ", "- [x] ", "- [X] ")):
                checked = s[3] in ("x", "X")
                self._rich_para(doc, ("☑ " if checked else "☐ ") + s[6:], style="List Bullet")
            elif s.startswith(("- ", "* ")):
                self._rich_para(doc, s[2:], style="List Bullet")
            else:
                self._rich_para(doc, s)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _rich_para(doc, text: str, style: str | None = None):
        """`**bold**` 인라인 마크다운을 굵게 반영해 문단 추가."""
        p = doc.add_paragraph(style=style)
        for i, part in enumerate(text.split("**")):
            if part:
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        return p
